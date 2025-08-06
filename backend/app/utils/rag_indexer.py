import os
import json
import logging
from pathlib import Path
from typing import Dict, List, Any
from datetime import datetime
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
from langchain_openai import AzureOpenAIEmbeddings
from ..config import logger, AZURE_CONFIG, output_dir
import PyPDF2
from docx import Document as DocxDocument

RAG_DIR = Path(output_dir) / "rag"
STANDARDS_RAG_DIR = Path(output_dir) / "standards-rag"

def get_embedding_client():
    """Initialize Azure OpenAI embedding client."""
    try:
        embed_model = AzureOpenAIEmbeddings(
            azure_endpoint=AZURE_CONFIG["AZURE_OPENAI_EMBED_API_ENDPOINT"],
            api_key=AZURE_CONFIG["AZURE_OPENAI_EMBED_API_KEY"],
            model=AZURE_CONFIG["AZURE_OPENAI_EMBED_MODEL"],
            azure_deployment=AZURE_CONFIG["AZURE_OPENAI_EMBED_DEPLOYMENT"],
            api_version=AZURE_CONFIG["AZURE_OPENAI_EMBED_VERSION"],
        )
        logger.info("Azure OpenAI embedding client initialized successfully")
        return embed_model
    except Exception as e:
        logger.error(f"Failed to initialize embedding client: {str(e)}")
        raise

embedding_client = get_embedding_client()

def test_embedding_service():
    """Test the embedding service."""
    try:
        test_text = "This is a test document for embedding."
        embedding = embedding_client.embed_query(test_text)
        logger.info(f"Embedding test successful. Dimension: {len(embedding)}")
        return True
    except Exception as e:
        logger.error(f"Embedding test failed: {str(e)}")
        return False

def extract_text_from_file(file_path: Path) -> str:
    """Extract text from PDF, DOC, DOCX, or TXT files."""
    logger.info(f"Extracting text from file: {file_path}")
    if file_path.suffix.lower() not in [".pdf", ".doc", ".docx", ".txt"]:
        logger.warning(f"Unsupported file type for {file_path}. Expected .pdf, .doc, .docx, or .txt.")
        return ""
    
    try:
        if file_path.suffix.lower() == ".pdf":
            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted
                return text
        elif file_path.suffix.lower() in [".doc", ".docx"]:
            doc = DocxDocument(file_path)
            return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        elif file_path.suffix.lower() == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        return ""

def index_standards_document(project_id: str, file_path: Path):
    """Index a standards document into a FAISS vector store."""
    logger.info(f"Indexing standards document for project: {project_id}, file: {file_path}")
    
    if not test_embedding_service():
        logger.error("Embedding service is not available")
        raise ValueError("Embedding service is not available")
    
    output_dir = STANDARDS_RAG_DIR / project_id / "faiss_index"
    output_dir.mkdir(parents=True, exist_ok=True)
    
    text_content = extract_text_from_file(file_path)
    if not text_content.strip():
        logger.warning(f"No content extracted from {file_path}")
        return
    
    document = Document(
        page_content=text_content,
        metadata={
            "source": file_path.name,
            "type": f"standards_{file_path.suffix.lstrip('.')}",
            "project_id": project_id
        }
    )
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    
    chunks = text_splitter.split_documents([document])
    logger.info(f"Split standards document into {len(chunks)} chunks")
    
    try:
        vector_store = FAISS.load_local(
            str(output_dir),
            embedding_client,
            allow_dangerous_deserialization=True
        )
        vector_store.add_documents(chunks)
        logger.info(f"Updated existing standards vector store for project: {project_id}")
    except Exception:
        vector_store = FAISS.from_documents(chunks, embedding_client)
        logger.info(f"Created new standards vector store for project: {project_id}")
    
    try:
        vector_store.save_local(str(output_dir))
        logger.info(f"Saved standards vector store to {output_dir}")
    except Exception as e:
        logger.error(f"Error saving standards vector store: {str(e)}")
        raise
    
    metadata_path = output_dir.parent / "metadata.json"
    metadata = {
        "project_id": project_id,
        "total_documents": 1,
        "total_chunks": len(chunks),
        "embedding_model": AZURE_CONFIG["AZURE_OPENAI_EMBED_MODEL"],
        "created_at": datetime.now().isoformat(),
    }
    
    with open(metadata_path, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, indent=2)
    
    logger.info(f"Standards document indexing completed for project: {project_id}")

def index_files_for_rag(project_id: str, cobol_json: Dict[str, Any], file_data: Dict[str, Any]):
    """Index COBOL files and analysis JSON for RAG."""
    logger.info(f"Indexing files for RAG: {project_id}")
    
    if not test_embedding_service():
        logger.error("Embedding service is not available")
        raise ValueError("Embedding service is not available")
    
    output_dir = RAG_DIR / project_id / "faiss_index"
    output_dir.mkdir(parents=True, exist_ok=True)
    
    documents = []
    
    # Process original file_data from request
    if file_data:
        logger.info(f"Found {len(file_data)} files in file_data to process")
        for file_name, file_info in file_data.items():
            content = file_info.get("content", "")
            file_type = file_info.get("type", "Unknown")
            if content:
                documents.append(Document(
                    page_content=f"File: {file_name}\nType: {file_type}\nContent:\n{content}",
                    metadata={
                        "source": file_name,
                        "type": f"cobol_{file_type.lower()}",
                        "project_id": project_id
                    }
                ))
                logger.info(f"Added document for file: {file_name}")
    
    # Process cobol_analysis.json
    analysis_path = Path(output_dir).parent.parent / "analysis" / project_id / "cobol_analysis.json"
    if analysis_path.exists():
        with open(analysis_path, "r", encoding="utf-8") as f:
            analysis_content = json.load(f)
            documents.append(Document(
                page_content=f"File: cobol_analysis.json\nType: Analysis\nContent:\n{json.dumps(analysis_content, indent=2)}",
                metadata={
                    "source": "cobol_analysis.json",
                    "type": "cobol_analysis",
                    "project_id": project_id
                }
            ))
            logger.info(f"Added document for cobol_analysis.json")
    
    if not documents:
        logger.error("No documents found to index")
        raise ValueError("No content found in COBOL analysis to index")
    
    logger.info(f"Prepared {len(documents)} documents for indexing")
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    
    chunks = text_splitter.split_documents(documents)
    logger.info(f"Split into {len(chunks)} chunks")
    
    try:
        vector_store = FAISS.load_local(
            str(output_dir),
            embedding_client,
            allow_dangerous_deserialization=True
        )
        vector_store.add_documents(chunks)
        logger.info(f"Updated existing vector store for project: {project_id}")
    except Exception:
        vector_store = FAISS.from_documents(chunks, embedding_client)
        logger.info(f"Created new vector store for project: {project_id}")
    
    try:
        vector_store.save_local(str(output_dir))
        logger.info(f"Saved vector store to {output_dir}")
    except Exception as e:
        logger.error(f"Error saving vector store: {str(e)}")
        raise
    
    metadata_path = output_dir.parent / "metadata.json"
    metadata = {
        "project_id": project_id,
        "total_documents": len(documents),
        "total_chunks": len(chunks),
        "embedding_model": AZURE_CONFIG["AZURE_OPENAI_EMBED_MODEL"],
        "created_at": datetime.now().isoformat(),
    }
    
    with open(metadata_path, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, indent=2)
    
    logger.info(f"RAG indexing completed successfully for project: {project_id}")

def load_vector_store(project_id: str):
    """Load and combine vector stores from both rag and standards-rag directories."""
    try:
        cobol_faiss_path = RAG_DIR / project_id / "faiss_index"
        standards_faiss_path = STANDARDS_RAG_DIR / project_id / "faiss_index"
        
        vector_stores = []
        
        if cobol_faiss_path.exists():
            cobol_vector_store = FAISS.load_local(
                str(cobol_faiss_path), 
                embedding_client, 
                allow_dangerous_deserialization=True
            )
            vector_stores.append(cobol_vector_store)
            logger.info(f"COBOL vector store loaded successfully for project: {project_id}")
        
        if standards_faiss_path.exists():
            standards_vector_store = FAISS.load_local(
                str(standards_faiss_path),
                embedding_client,
                allow_dangerous_deserialization=True
            )
            vector_stores.append(standards_vector_store)
            logger.info(f"Standards vector store loaded successfully for project: {project_id}")
        
        if not vector_stores:
            logger.warning(f"No vector stores found for project: {project_id}")
            return None
        
        if len(vector_stores) > 1:
            combined_vector_store = vector_stores[0]
            for vs in vector_stores[1:]:
                combined_vector_store.merge_from(vs)
            logger.info(f"Combined {len(vector_stores)} vector stores for project: {project_id}")
            return combined_vector_store
        return vector_stores[0]
        
    except Exception as e:
        logger.error(f"Error loading vector store for project {project_id}: {str(e)}")
        return None

def query_vector_store(vector_store, query: str, k: int = 3):
    """Query the combined vector store with similarity search."""
    try:
        if not vector_store:
            logger.warning("Vector store is None")
            return []
        
        logger.info(f"Performing similarity search with query: '{query}' and k={k}")
        results = vector_store.similarity_search(query, k=k)
        logger.info(f"Found {len(results)} results")
        
        for i, result in enumerate(results):
            logger.info(f"Result {i+1}: source={result.metadata.get('source', 'unknown')}, type={result.metadata.get('type', 'unknown')}")
            logger.info(f"Result {i+1} content preview: {result.page_content[:100]}...")
        
        return results
        
    except Exception as e:
        logger.error(f"Error querying vector store: {str(e)}")
        return []